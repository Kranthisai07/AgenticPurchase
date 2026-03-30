# ==============================================================
#  📘 Generate Methodology Section for Agentic Purchase System
#  Output: Agentic_Purchase_Methodology_Report.docx
# ==============================================================

from docx import Document

doc = Document()
doc.add_heading("Implementation and Methodology – Agentic Purchase Research System", level=1)

doc.add_paragraph(
    "The Agentic Purchase Research System was implemented as a modular, research-grade multi-agent framework integrating computer vision, reasoning, trust validation, and simulated checkout under a Coordinator–Worker (Saga) architecture. Each component was developed as an independent FastAPI microservice communicating over HTTP, while a central Coordinator maintained the overall saga state, timeouts, compensations, and token budgeting."
)

# --------------------------------------------------------------
# System Architecture
# --------------------------------------------------------------
doc.add_heading("System Architecture", level=2)
doc.add_paragraph(
    "The backend comprises six FastAPI applications: one Coordinator and five stateless Worker Agents (Vision, Intent, Sourcing, Trust, Checkout). Each agent exposes independent routes, executes a specific role, and returns structured JSON adhering to shared Pydantic models."
)
architecture_points = [
    "Coordinator (Port 8000): Orchestrates the five agents, maintains saga state transitions (S1–S5), applies timeout and retry policies, and logs all events to JSONL.",
    "Vision Agent (8101): Uses the Google Cloud Vision API for object detection, OCR, and color analysis; optionally invokes an LLM to refine ambiguous hypotheses.",
    "Intent Agent (8102): Conversationally confirms product details (color, quantity, price range) using heuristic or LangChain-driven logic.",
    "Sourcing Agent (8103): Retrieves offers from a mock catalog (data/mock_catalog.json), applies weighted heuristics for scoring, and can rerank top results with LLM support.",
    "Trust Agent (8104): Validates vendor credibility via TLS, WHOIS, and policy heuristics; medium/high risk triggers Coordinator compensation routines.",
    "Checkout Agent (8105): Simulates secure payment (Luhn validation, CVV/expiry check), ensuring idempotent checkout through unique request keys."
]
for item in architecture_points:
    doc.add_paragraph(item, style="List Bullet")

# --------------------------------------------------------------
# Tools and Frameworks
# --------------------------------------------------------------
doc.add_heading("Tools and Frameworks", level=2)
tools = [
    "FastAPI (Python 3.10) for asynchronous backend services.",
    "LangChain for optional LLM orchestration (OpenAI, Gemini, Ollama).",
    "Google Cloud Vision API for image recognition and OCR.",
    "Pydantic for schema validation and inter-agent contracts.",
    "HTTPX / Uvicorn for async communication.",
    "React + Tailwind (Vite) for the chat UI with image upload, quick actions, and voice input.",
    "Pandas + JSONL for offline metric parsing and evaluation."
]
for tool in tools:
    doc.add_paragraph(tool, style="List Bullet")

doc.add_paragraph(
    "All services are container-ready and launched via a PowerShell script (scripts/run-all.ps1) that spawns six console windows for concurrent operation."
)

# --------------------------------------------------------------
# System Connectivity
# --------------------------------------------------------------
doc.add_heading("System Connectivity", level=2)
doc.add_paragraph(
    "The Coordinator communicates with all agents over RESTful HTTP, configured through environment variables:"
)
doc.add_paragraph(
    "AGENT_VISION_URL = http://127.0.0.1:8101\n"
    "AGENT_INTENT_URL = http://127.0.0.1:8102\n"
    "AGENT_CHECKOUT_URL = http://127.0.0.1:8105"
)
doc.add_paragraph(
    "Each agent consumes and produces JSON conforming to shared schemas (libs/schemas/models.py) such as ProductHypothesis, PurchaseIntent, Offer, TrustAssessment, and Receipt. The Coordinator passes each state’s output as the next state’s input, forming the sequential saga chain S1–S5. For reliability, headers such as X-Request-ID and Idempotency-Key propagate through all HTTP calls. In local development mode, the Coordinator can import agent modules directly instead of HTTP calls to minimize latency."
)

# --------------------------------------------------------------
# Implementation Details
# --------------------------------------------------------------
doc.add_heading("Implementation Details", level=2)
details = [
    "Coordinator Logic: Defined in apps/coordinator/saga.py, using asynchronous event loops with timeout enforcement (asyncio.wait_for). It monitors token budgets and logs metrics through metrics_tokens.py.",
    "Vision Pipeline: Captures image input, sends it to Google Cloud Vision, and constructs structured hypotheses with labels, brands, and confidence scores.",
    "Intent Processing: Parses user input or LLM-generated dialogue responses into structured purchase intents.",
    "Offer Sourcing: Uses deterministic heuristics for scoring (price, shipping time, match quality) and optionally invokes reranking LLMs with enforced token budgets.",
    "Trust Validation: Performs heuristic checks on domains and TLS configurations; outputs a risk level that determines whether compensation is required.",
    "Checkout Simulation: Ensures payment validation and idempotent transaction safety. If trust risk is flagged, the Coordinator retries with a secondary offer."
]
for d in details:
    doc.add_paragraph(d, style="List Bullet")

# --------------------------------------------------------------
# System Workings
# --------------------------------------------------------------
doc.add_heading("System Workings", level=2)
steps = [
    "S1 Capture (Vision): The UI uploads an image; Agent 1 identifies the product and returns a structured ProductHypothesis.",
    "S2 Intent Confirmation: Agent 2 confirms user constraints like color and budget, producing a PurchaseIntent.",
    "S3 Sourcing: Agent 3 scores available offers and optionally performs LLM reranking under a 1.5k token cap.",
    "S4 Trust Evaluation: Agent 4 assigns a TrustAssessment. If risk ≥ medium, the Coordinator triggers a fallback to a safer vendor.",
    "S5 Checkout: The user confirms; Agent 5 validates payment and returns a final Receipt object."
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

doc.add_paragraph(
    "Each saga run is logged as JSONL, containing state names, timestamps, retry counts, and token usage metrics."
)

# --------------------------------------------------------------
# Data Flow and Logging
# --------------------------------------------------------------
doc.add_heading("Data Flow and Logging", level=2)
doc.add_paragraph(
    "All saga executions are tracked within logs/eval.log. Each event includes state ID, timestamp, token estimate, agent status, and latency. The evaluation script eval_report.py aggregates results into CSVs (eval_summary.csv, token_summary.csv) for later analysis."
)

# --------------------------------------------------------------
# Summary
# --------------------------------------------------------------
doc.add_heading("Summary", level=2)
doc.add_paragraph(
    "This methodology demonstrates the design of a distributed, token-aware multi-agent purchase pipeline that balances autonomy and control. Its modular structure supports empirical evaluation of latency, accuracy, trust precision, and token efficiency for future research on agentic AI commerce systems."
)

# Save the document
file_path = "Agentic_Purchase_Methodology_Report.docx"
doc.save(file_path)

print(f"\n✅ File generated successfully! Saved as: {file_path}")
